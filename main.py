# 1. Importações
import fitz  # PyMuPDF
import docx
import re
from docx import Document
from PyPDF2 import PdfReader
from tkinter import Tk, filedialog

# 2. Funções (extrair_referencias_completas, identificar_cor, etc.)
def extrair_referencias_completas(pdf_path, pagina_inicial):
    """Extrai a seção de referências do PDF e ajusta a numeração das páginas."""
    palavras_chave = ["References", "Referências", "Bibliography", "Works Cited", "Cited References"]
    possivel_fim = ["Appendix", "Acknowledgments", "Acknowledgements", "Resumo", "Abstract", "Conclusão"]

    referencias_extraidas = []
    capturando = False

    with fitz.open(pdf_path) as doc:
        for i, pagina in enumerate(doc, start=pagina_inicial):
            texto = pagina.get_text("text")
            linhas = texto.split("\n")

            for linha in linhas:
                linha_limpa = linha.strip()

                if any(re.match(rf"^\s*{palavra}\s*$", linha_limpa, re.IGNORECASE) for palavra in palavras_chave):
                    capturando = True
                    referencias_extraidas.append(f"[Página {i}] {linha_limpa}")

                if capturando:
                    referencias_extraidas.append(linha_limpa)

                if any(re.match(rf"^\s*{fim}\s*$", linha_limpa, re.IGNORECASE) for fim in possivel_fim) and capturando:
                    capturando = False
                    break

    return referencias_extraidas if referencias_extraidas else ["Nenhuma referência bibliográfica encontrada."]
def identificar_cor(rgb):
    cores = {
        "amarelo": (255, 237, 0),
        "azul": (45, 98, 244),
        "verde": (51, 158, 0),
        "vermelho": (255, 0, 0)
    }
    limite = 75
    cor_mais_proxima = min(
        cores.keys(),
        key=lambda cor: sum((a - b) ** 2 for a, b in zip(rgb, cores[cor]))
    )
    distancia = sum((a - b) ** 2 for a, b in zip(rgb, cores[cor_mais_proxima])) ** 0.5

    return cor_mais_proxima if distancia <= limite else "Outros"
def extrair_texto_por_posicao(pdf_path, pagina_inicial):
    doc = fitz.open(pdf_path)
    total_paginas = len(doc)
    elementos = []

    for i, page in enumerate(doc, start=pagina_inicial):
        palavras = page.get_text("words")
        if not page.annots():
            continue

        for annot in page.annots():
            if annot.type[0] != 8:
                continue  # apenas highlight

            cor_base = annot.colors.get("stroke") or annot.colors.get("fill") or (0, 0, 0)
            cor_rgb = tuple(int(c * 255) for c in cor_base)
            cor_nome = identificar_cor(cor_rgb)

            quad_points = annot.vertices
            quads = [fitz.Quad(quad_points[i:i+4]) for i in range(0, len(quad_points), 4)]
            palavras_selecionadas = []

            for quad in quads:
                for w in palavras:
                    if quad.rect.intersects(fitz.Rect(w[:4])):
                        palavras_selecionadas.append((w[0], w[1], w[4]))

            if palavras_selecionadas:
                palavras_selecionadas.sort(key=lambda w: (w[1], w[0]))
                texto = " ".join(w[2] for w in palavras_selecionadas).strip()
                y_medio = sum(w[1] for w in palavras_selecionadas) / len(palavras_selecionadas)

                elementos.append({
                    "pagina": i,
                    "y": y_medio,
                    "cor": cor_nome,
                    "texto": texto
                })

    elementos.sort(key=lambda e: (e["pagina"], e["y"]))

    titulos = [el for el in elementos if el["cor"] == "vermelho"]
    grifos = [el for el in elementos if el["cor"] == "amarelo"]

    secoes = {}
    ordem_secao = []

    for idx, titulo in enumerate(titulos):
        pag_ini = titulo["pagina"]
        y_ini = titulo["y"]

        if idx + 1 < len(titulos):
            proximo_titulo = titulos[idx + 1]
            pagina_limite = proximo_titulo["pagina"]
            y_limite = proximo_titulo["y"]
        else:
            # Não há próximo título: determinar até onde os grifos vão
            grifos_na_secao = [
                g for g in grifos
                if g["pagina"] > pag_ini or (g["pagina"] == pag_ini and g["y"] >= y_ini)
            ]
            if grifos_na_secao:
                ultimo_grifo = max(grifos_na_secao, key=lambda g: (g["pagina"], g["y"]))
                pagina_limite = ultimo_grifo["pagina"]
                y_limite = ultimo_grifo["y"]
            else:
                pagina_limite = pag_ini
                y_limite = float('inf')  # Não tem grifos, pega o título sozinho

        grifos_da_secao = [
            g["texto"]
            for g in grifos
            if (g["pagina"] > pag_ini or (g["pagina"] == pag_ini and g["y"] >= y_ini))
            and (g["pagina"] < pagina_limite or (g["pagina"] == pagina_limite and g["y"] <= y_limite))
        ]

        if grifos_da_secao:
            nome_secao = f"(: {pag_ini} - {pagina_limite}) {titulo['texto']}"
            secoes[nome_secao] = grifos_da_secao
            ordem_secao.append(nome_secao)

    secoes_ordenadas = {secao: secoes[secao] for secao in ordem_secao}
    return secoes_ordenadas


# Substituir a função upload_arquivo para usar filedialog
def upload_arquivo():
    root = Tk()
    root.withdraw()
    caminho_pdf = filedialog.askopenfilename(
        title="Selecione o arquivo PDF",
        filetypes=[("Arquivos PDF", "*.pdf")]
    )
    return caminho_pdf

# Substituir salvar_extrato (sem files.download)
def salvar_extrato(nome_extrato, secoes):
    doc = Document()
    doc.add_heading(nome_extrato, level=1)

    for secao, trechos in secoes.items():
        doc.add_heading(secao, level=2)
        for trecho in trechos:
            doc.add_paragraph(f"{trecho}")

    nome_arquivo = f"{nome_extrato}.docx"
    doc.save(nome_arquivo)
    print(f"\n✅ Documento salvo como '{nome_arquivo}'.")

# 3. Execução principal
if __name__ == "__main__":
    pdf_path = upload_arquivo()
    pagina_inicial = int(input("\nDigite o número da página inicial: "))
    nome_extrato = input("\nDigite o nome do extrato (sem extensão): ").strip() or "Extrato"

    print("\n🖍️ Extraindo referências bibliográficas...")
    referencias_extraidas = extrair_referencias_completas(pdf_path, pagina_inicial)

    print("\n🖍️ Extraindo trechos grifados por cor...")
    textos_por_topico = extrair_texto_por_posicao(pdf_path, pagina_inicial)

    print("\n🖍️ Organizando os trechos em seções...")
    secoes = {**textos_por_topico, "Referências Bibliográficas": referencias_extraidas}

    print("\n💾 Salvando o extrato no formato .docx...")
    salvar_extrato(nome_extrato, secoes)
